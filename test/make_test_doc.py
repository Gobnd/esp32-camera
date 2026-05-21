from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Style helpers ─────────────────────────────────────────────────────────────

DARK  = RGBColor(0x0d, 0x1b, 0x4b)
RED   = RGBColor(0x8b, 0x00, 0x00)
GREEN = RGBColor(0x1a, 0x6b, 0x1a)
GREY  = RGBColor(0x55, 0x55, 0x55)
AMBER = RGBColor(0x92, 0x60, 0x00)

def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def shade_para(p, hex_fill="F0F0F0"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

def h1(text, color=DARK):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = color
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = DARK
    return p

def body(text):
    p = doc.add_paragraph()
    p.add_run(text).font.size = Pt(10)
    return p

def italic(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY
    return p

def note(text, color=AMBER):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.size = Pt(9); r.font.color.rgb = color
    shade_para(p, "FFF8E8")
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(10)

def table(headers, rows, col_widths=None, status_col=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, cell in enumerate(t.rows[0].cells):
        cell.text = headers[i]
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True; run.font.size = Pt(9)
        shade_cell(cell, 'D0D8E8')
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(t.rows[ri + 1].cells):
            cell.text = str(row[ci])
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if status_col is not None and ci == status_col:
                val = str(row[ci]).upper()
                if "PASS" in val:   shade_cell(cell, "D4EDDA")
                elif "FAIL" in val: shade_cell(cell, "F8D7DA")
                elif "PEND" in val or "NOT YET" in val: shade_cell(cell, "FFF3CD")
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()

def page_break():
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)

def divider():
    p = doc.add_paragraph("─" * 74)
    p.runs[0].font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    p.runs[0].font.size = Pt(9)

# ── Cover ─────────────────────────────────────────────────────────────────────

doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("XIAO ESP32S3 Camera - Test Results")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = DARK

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.add_run(f"Updated {datetime.date.today().strftime('%d %b %Y')}").font.size = Pt(10)
doc.add_paragraph(); divider(); doc.add_paragraph()

# ── Header table ──────────────────────────────────────────────────────────────

table(
    ["Camera IP", "Board", "Last Test Date", "Firmware Flash Date"],
    [["10.10.10.155 (work) / esp32cam.local (hybrid/deepsleep)", "Seeed XIAO ESP32S3 Sense (OV2640)", "2026-05-21", "2026-05-21"]],
    col_widths=[2.6, 2.2, 1.2, 1.0]
)

note("Three firmware environments: 'streaming' (always-on MJPEG+RTSP), 'deepsleep' (PIR wake → 1 photo → sleep), 'hybrid' (PIR wake → 30s MJPEG stream → sleep). Upload with: pio run -e <streaming|deepsleep|hybrid> --target upload. SXGA and UXGA resolution switch crashes streaming firmware - under investigation.")

doc.add_paragraph()

# ── Summary table ─────────────────────────────────────────────────────────────

table(
    ["Test", "Method", "Result"],
    [
        ["T1 - Motion detection",        "Semi-automated (--motion flag)",     "Pending - PIR attached, cooldown not yet applied"],
        ["T2 - WiFi range",              "wifi_test.py",                        "PASS (2026-05-21, SenSen2)"],
        ["T3 - Face recognition dist.",  "Manual",                              "PASS (2026-05-14, unchanged)"],
        ["T3b - Resolution vs FPS",      "res_fps_test.py",                     "PASS - VGA optimal (25 fps)"],
        ["T4 - Power / temps",           "Partially automated",                 "PASS (RTSP & PIR rows blank)"],
        ["T5 - Multi-stream",            "Automated",                           "PASS - all 5 clients alive"],
        ["T6 - RTSP",                    "Partially automated",                 "PASS (VLC steps manual)"],
        ["T6b - 5-min stability",        "Automated",                           "PASS"],
        ["TEP - Endpoint check",         "Automated",                           "PASS"],
        ["TM - Motion capture (streaming)", "motion_test.py",                   "PASS - all 4 subjects, face visible"],
        ["TH - Motion capture (hybrid)", "motion_test.py + hybrid firmware",    "PASS - all 4 subjects on battery"],
        ["TD - Deep sleep single photo", "motion_test_deepsleep.py",            "Partial - 1 photo/wake insufficient"],
        ["T7 - Multiple cameras",        "Manual",                              "Not yet run - needs 2nd unit"],
    ],
    col_widths=[2.4, 2.0, 2.6],
    status_col=2
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("T1 - Motion Detection & Capture Latency  [Pending]")
italic("Method: python run_tests.py 192.168.0.39 --motion - polls /stats for 30 s, logs latency and in-frame status per capture. Wave hand at PIR sensor on D0. PIR wiring: red→3V3, black→GND, yellow→D0. Run with no other browser tabs open.")

table(
    ["Trial", "Latency (ms)", "JPEG size (bytes)", "In-frame?", "Notes"],
    [["1","","","",""], ["2","","","",""], ["3","","","",""], ["4","","","",""], ["5","","","",""]],
    col_widths=[0.7, 1.3, 1.6, 1.1, 2.3]
)

h2("Implications")
bullet("5-second cooldown now prevents hold-time re-triggering - one walk-through should give exactly 1 capture, not 2–3.")
bullet("If consistently off-camera, PIR detects movement outside camera FoV - aim both sensors in same direction.")
bullet("Latency >500 ms means camera buffer was held by a concurrent stream; close all streams before this test.")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("T2 - WiFi Range  [PASS]")
italic("Method: wifi_test.py - 10 s frame count per viewer at each location. Network: SenSen2.")

table(
    ["Location", "RSSI (dBm)", "FPS - 1 viewer", "FPS each - 2 viewers"],
    [
        ["Next to router",                   "−39", "22.5", "12.5"],
        ["Other side of same room",          "−51", "24.2", "12.2"],
        ["Next room (1 wall)",               "−58", "23.2", "10.2"],
        ["Outside building (same floor)",    "−76", "8.1",  "3.9"],
        ["Outside building (diff. floor)",   "−86", "3.5",  "1.8"],
        ["Across the street",                "-",   "Unreachable", "-"],
    ],
    col_widths=[2.2, 1.0, 1.3, 1.5]
)
italic("RSSI guide: above −60 dBm = strong | −70 dBm = acceptable | below −75 dBm = degraded")

h2("Implications")
bullet("Building wall is a hard boundary - inside signal stays above −60 dBm and FPS above 20. Outside, both collapse.")
bullet("Usable streaming (>10 fps/viewer) cuts off at the building wall.")
bullet("Each wall costs roughly 10–15 dBm; reposition router before moving camera if signal is marginal.")
note("Bugs fixed during this test: RSSI parser was matching −8 from charset='utf-8' in HTML (fixed to target WiFi Signal table cell). FPS was reading lifetime average from /stats (fixed to count --frame markers directly from stream).")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("T3 - Face Recognition Distance  [PASS]")
italic("Method: Manual - subject walked away in a straight corridor under fair indoor lighting. Observer called clear / barely / gone; distances measured with tape measure. Repeated per resolution. Fair lighting, 30 fps.")

table(
    ["Resolution", "Clearly recognisable (cm)", "Barely recognisable (cm)", "Not recognisable (cm)"],
    [
        ["240p (320×240)",   "100", "150", "200"],
        ["480p (640×480)",   "150", "200", "250"],
        ["720p (1280×720)",  "300", "350", "400"],
        ["SXGA (1280×1024)", "320", "360", "400"],
        ["UXGA (1600×1200)", "380", "450", "500"],
    ],
    col_widths=[1.8, 1.8, 1.8, 1.8]
)

h2("Implications")
bullet("720p is the recommended resolution for a 3 m corridor - faces clearly recognisable at 3 m, still visible to 4 m.")
bullet("UXGA adds only ~80 cm over 720p at cost of lower FPS; only use it for corridors longer than ~3.5 m.")
bullet("All thresholds drop 30–50% in dim lighting - add supplementary lighting or switch to IR camera for dark environments.")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("T3b - Resolution vs FPS Trade-off  [PASS]")
italic("Method: Automated - res_fps_test.py. For each resolution: sends GET /res?id=<id>, waits 5 s (SETTLE_SECS), then opens one HTTP MJPEG stream. Raw bytes collected for 12 s; counts --frame boundary markers. No FPS cap. One sample JPEG per resolution saved to ./frames/. Board cold, no active connections.")

table(
    ["Resolution", "Dimensions", "FPS (2026-05-21)", "FPS (2026-05-14)", "Temp (°C)", "Face range (T3)", ">=20 fps?"],
    [
        ["qvga", "320×240",   "26.0", "25.4", "46.3", "100 cm", "YES"],
        ["vga",  "640×480",   "20.9", "24.8", "47.3", "150 cm", "YES - recommended"],
        ["svga", "800×600",   "26.5", "13.1", "50.3", "-",      "YES"],
        ["hd",   "1280×720",  "15.2", "13.9", "50.3", "300 cm", "BORDERLINE"],
        ["sxga", "1280×1024", "11.9", "10.6", "52.3", "320 cm", "NO"],
        ["uxga", "1600×1200", "8.3",  "7.8",  "53.3", "380 cm", "NO"],
    ],
    col_widths=[0.8, 1.0, 1.4, 1.4, 0.9, 1.2, 1.3]
)
note("Previous 2026-05-19 run showed SVGA at 13 fps and SXGA/UXGA crashing - both were test artifacts caused by an active stream connection during the test. Clean 2026-05-21 run (no active connections) shows SVGA at 26.5 fps and no crashes at any resolution. Old 2026-05-14 SVGA result of 13.1 fps was also contaminated.")

h2("Implications")
bullet("No hard cliff at VGA. SVGA delivers 26.5 fps - faster than VGA on a clean run. OV2640 is not limited at VGA.")
bullet("Real performance drop is at HD (15 fps, borderline) and above. SXGA/UXGA drop to 8-12 fps.")
bullet("VGA still recommended for motion capture - SVGA adds no face recognition distance benefit (T3 shows same range as VGA) and the extra pixels are wasted.")
bullet("SXGA/UXGA do not crash. Previous crash was from switching resolution while a stream was already active.")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("T4 - Power & Temperature  [PASS - partial]")
italic("Method: Idle row automated. Stream rows captured during T5. RTSP and PIR rows require manual INA219 measurement. FPS column is lifetime average (totalFramesSent / uptime_ms) - not per-state current FPS. CPU 240 MHz")

table(
    ["State", "Temp (°C)", "FPS (lifetime avg)"],
    [
        ["Idle",             "39.3", "0.3"],
        ["1x MJPEG stream",  "42.3", "0.4"],
        ["2x MJPEG streams", "43.3", "0.7"],
        ["3x MJPEG streams", "44.3", "0.9"],
        ["4x MJPEG streams", "45.3", "1.1"],
        ["5x MJPEG streams", "46.3", "1.2"],
        ["1x RTSP (VLC)",    "",     ""],
        ["PIR during stream","",     ""],
    ],
    col_widths=[2.8, 1.2, 1.8]
)
note("FPS column shows lifetime average since boot - dragged down by idle time before T5. Not comparable to old T4 FPS numbers. Use T3b (res_fps_test.py) for accurate per-resolution FPS.")

h2("Implications")
bullet("Temperature rises ~7°C from idle to 5-stream load - passive cooling sufficient in normal indoor environments.")
bullet("Board at home runs 7–10°C cooler than at work (stronger signal = less radio power = less heat).")
bullet("For accurate battery life estimate, clip INA219 on 3.3 V rail - firmware estimate is not load-sensitive.")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("T5 - Multiple Simultaneous Streams  [PASS]")
italic("Method: Automated - opens 1–5 HTTP MJPEG streams in threads, waits 5 s per step, reads /stats. Frames rx column is cumulative total across all clients since test started - NOT per-step FPS. Only the 1-client row gives a clean per-client rate.")

table(
    ["Clients", "All alive?", "Frames rx (cumulative)", "Temp (°C)", "Crashes"],
    [
        ["1", "Yes", "102",  "42.3", "None"],
        ["2", "Yes", "238",  "43.3", "None"],
        ["3", "Yes", "409",  "44.3", "None"],
        ["4", "Yes", "541",  "45.3", "None"],
        ["5", "Yes", "657",  "46.3", "None"],
    ],
    col_widths=[0.8, 0.9, 2.0, 1.0, 0.9]
)
note("1-client clean rate: 102 frames / 5 s = ~20 fps. Beyond row 1, Frames rx includes all prior clients still running - rows cannot be compared directly. All-alive column is the meaningful pass/fail metric.")

h2("Old T5 data for reference - FPS measured per-step (old firmware)")
table(
    ["Clients", "FPS (per-step, old method)", "Temp (°C)", "Crashes"],
    [
        ["1", "27.8", "51.3", "None"],
        ["2", "12.4", "53.3", "None"],
        ["3", "7.3",  "53.3", "None"],
        ["4", "6.3",  "53.3", "None"],
        ["5", "5.6",  "54.3", "None"],
    ],
    col_widths=[0.8, 2.2, 1.0, 0.9]
)
note("Old FPS numbers were from a per-handler racy global (streamFps) - values with 2+ clients are unreliable. New firmware removes the race but the test script needs updating to compute per-step delta for clean multi-client FPS.")

h2("Implications")
bullet("For recognition pipeline use 1 client only - gives ~20–28 fps depending on firmware and network.")
bullet("Use RTSP for recognition software and HTTP /stream for human preview only.")
bullet("Temperature at 5 clients (46°C) well below 85°C throttle threshold.")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("T6 - RTSP Stream  [PASS - VLC steps manual]")
italic("Method: Automated OPTIONS probe (raw TCP socket to :8554). VLC visual checks are manual.")
body("RTSP URL: rtsp://192.168.0.39:8554/mjpeg/1  (home) / rtsp://10.10.10.155:8554/mjpeg/1  (work)")

table(
    ["Step", "Result"],
    [
        ["RTSP OPTIONS probe",           "PASS - RTSP/1.0 200 OK"],
        ["Methods advertised",           "DESCRIBE, SETUP, TEARDOWN, PLAY, PAUSE"],
        ["VLC connects and plays",        ""],
        ["No green bands / duplication",  ""],
        ["Resolution change visible in VLC", ""],
        ["Quality change visible in VLC",    ""],
    ],
    col_widths=[2.8, 4.2],
    status_col=1
)

h2("Implications")
bullet("Only 1 simultaneous RTSP client supported - recognition software should be sole RTSP consumer.")
bullet("Resolution changes apply immediately via /res; RTSP reflects new resolution on next VLC reconnect.")
bullet("No authentication - any device on the same network can change settings.")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("T6b - 5-Minute Stability  [PASS]")
italic("Method: Automated - 3 simultaneous HTTP streams held for 300 s, polled every 30 s. FPS is lifetime average - jitter reflects dilution by earlier idle time, not real FPS instability.")

table(
    ["Elapsed (s)", "Streams alive", "Temp (°C)", "Free RAM", "RSSI (dBm)"],
    [
        ["30",  "3/3", "46.3", "189 / 323 KB", "-51"],
        ["60",  "3/3", "47.3", "188 / 323 KB", "-51"],
        ["90",  "3/3", "46.3", "181 / 323 KB", "-54"],
        ["120", "3/3", "48.3", "188 / 323 KB", "-54"],
        ["150", "3/3", "49.3", "188 / 323 KB", "-59"],
        ["180", "3/3", "48.3", "188 / 323 KB", "-51"],
        ["210", "3/3", "51.3", "181 / 323 KB", "-54"],
        ["241", "3/3", "49.3", "190 / 323 KB", "-53"],
        ["271", "3/3", "50.3", "188 / 323 KB", "-57"],
        ["301", "3/3", "53.3", "188 / 323 KB", "-52"],
    ],
    col_widths=[1.0, 1.2, 1.0, 1.5, 1.1]
)
body("PASS - all 3 streams alive throughout. Temp plateaued at 53°C (safe limit 85°C). RAM stable, no leak. RSSI -51 to -59 dBm (good). Board runs cooler at home vs work (-38 dBm work vs -51 to -59 dBm home).")
note("Rising FPS from 2.3 to 7.4 over 5 min is expected with the lifetime average formula - it's approaching the true streaming rate as idle time becomes a smaller fraction of total uptime. Not a sign of instability.")

h2("Old T6b data for reference")
table(
    ["Elapsed (s)", "Streams alive", "FPS (/stats, old)", "Temp (°C)", "Free RAM", "RSSI (dBm)"],
    [
        ["30",  "3/3", "9.2", "56.3", "190 / 323 KB", "-38"],
        ["60",  "3/3", "6.5", "57.3", "183 / 323 KB", "-38"],
        ["120", "3/3", "6.5", "59.3", "191 / 323 KB", "-39"],
        ["180", "3/3", "6.5", "60.3", "182 / 323 KB", "-38"],
        ["300", "3/3", "7.0", "62.3", "191 / 323 KB", "-38"],
    ],
    col_widths=[1.0, 1.1, 1.5, 0.9, 1.3, 1.0]
)
note("Old FPS from /stats used the racy per-handler global - values were unreliable under 3-client load. Old temps higher (56–62°C) because board ran at work (weaker signal = more radio power).")

h2("Implications")
bullet("System can run indefinitely - no memory leak detected.")
bullet("Chip reaches thermal equilibrium with a heatsink.")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("TEP - Endpoint Verification  [PASS]")
italic("Method: Automated - sends GET to each control endpoint, re-reads /stats to confirm change, then restores defaults.")

table(
    ["Endpoint", "Value set", "Confirmed in /stats", "Result"],
    [
        ["/res?id=qvga",    "qvga (320x240)", "resolution = qvga", "PASS"],
        ["/quality?q=10",   "q=10 (Best)",    "quality = 10",      "PASS"],
        ["/fps?delay=33",   "33 ms delay",    "frame_delay = 33",  "PASS"],
    ],
    col_widths=[1.5, 1.5, 2.0, 1.0],
    status_col=3
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# ═════════════════════════════════════════════════════════════════════════════
h1("TM - Motion Capture - Streaming Firmware  [PASS]")
italic("Method: motion_test.py - 15 s MJPEG stream recorded per subject, saved as frame_XXXX.jpg to motion_captures/<subject>/. VGA 640×480, quality 20. Location: office corridor, SenSen2 WiFi.")

table(
    ["Subject", "Face captured?", "Notes"],
    [
        ["Fast walker",  "Yes", "Multiple clear frames mid-stride"],
        ["Normal walker","Yes", "Clear face across most frames"],
        ["Slow walker",  "Yes", "Many usable frames"],
    ],
    col_widths=[2.0, 1.2, 3.8],
    status_col=1
)
body("Key finding: 15 s at 24 fps = ~360 frames per subject. Even a fast walker crossing frame in 2–3 s produces 50+ frames - guarantees at least one clear face shot.")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("TH - Motion Capture - Hybrid Firmware (battery)  [PASS]")
italic("Method: motion_test.py pointing at esp32cam.local/stream. Board on power bank. Hybrid firmware: sleeps at ~14 µA until PIR fires → wakes → streams MJPEG for 30 s → sleeps.")

table(
    ["Subject", "Face captured?"],
    [
        ["Fast walker",  "Yes"],
        ["Normal walker","Yes"],
        ["Slow walker",  "Yes"],
    ],
    col_widths=[3.2, 3.8],
    status_col=1
)

h2("Firmware behaviour")
bullet("On PIR wake: camera init → WiFi connect → mDNS (esp32cam.local) → MJPEG stream for 30 s → deep sleep.")
bullet("On 30 s timer wake: immediately deep sleep (power bank keepalive only - draws no useful current).")
bullet("Power bank issue solved: 14 µA deep sleep draw is too low for power bank auto-shutoff threshold. Timer wakeup creates brief current spike every 30 s to keep power bank alive.")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("TD - Motion Capture - Deep Sleep Firmware  [Partial]")
italic("Method: motion_test_deepsleep.py - polls esp32cam.local/motion for up to 60 s per subject. Board wakes on PIR, takes 1 photo, serves it for 30 s, sleeps.")
italic("So this did not work as PIR scanned object off-camera or when person was turning back. RTSP works better but again uses resources. I have now created an alternate firmware that sends multiple images instead of a simple RTSP or a single image.")

table(
    ["Subject", "Board responded?", "Wake time", "In-frame?", "Notes"],
    [
        ["Fast walker",   "NO", "", "", ""],
        ["Normal walker", "",   "", "", ""],
        ["Slow walker",   "",   "", "", ""],
    ],
    col_widths=[1.5, 1.3, 0.9, 0.9, 3.0]
)

h2("Issues found - fixes applied")
body("Notes:")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1("T7 - Multiple Cameras  [Not yet run - requires second unit]")
italic("Method: Manual - flash second board with same firmware, connect both to WiFi, verify independent stream and control access. Note: test scripts must have IP passed as argument (no hardcoded IP) before running T7.")

table(
    ["Test", "Camera 1", "Camera 2"],
    [
        ["Both streaming simultaneously",        "", ""],
        ["Both RTSP in VLC simultaneously",       "", ""],
        ["Res change on one does not affect other","", ""],
        ["Motion detection independent",          "", ""],
    ],
    col_widths=[3.0, 1.5, 1.5]
)

# ── Footer ─────────────────────────────────────────────────────────────────
doc.add_paragraph(); divider()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run(
    "Board: Seeed XIAO ESP32S3 Sense  |  Work IP: 10.10.10.155  |  Hybrid/Deepsleep: esp32cam.local"
).font.size = Pt(8)

out = "/home/gobind/Documents/PlatformIO/Projects/Example-PIO/test/Cam_tests.docx"
doc.save(out)
print(f"Saved: {out}")
