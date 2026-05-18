from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles ───────────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
    return p

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("Note: " + text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x6b, 0x1a)
    # shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def divider():
    doc.add_paragraph("─" * 72)

# ── Cover ────────────────────────────────────────────────────────────────────

doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ESP32 Camera + PIR Motion System")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(f"Study notes — generated {datetime.date.today().strftime('%d %b %Y')}").font.size = Pt(11)

doc.add_paragraph()
divider()
doc.add_paragraph()

# ── Section 1 ────────────────────────────────────────────────────────────────

h1("1. What the PIR sensor actually is")

body(
    "PIR stands for Passive Infrared. The word passive is the key — the sensor does not "
    "emit anything. It listens. Every warm body (human, animal) radiates infrared heat. "
    "When that heat source moves across the sensor's field of view, the IR level shifts "
    "and the sensor's internal chip trips a comparator and pulls its output pin HIGH."
)

body(
    "The sensor outputs a single digital signal: HIGH (3.3 V) = motion detected, "
    "LOW (0 V) = quiet. That single wire connects to pin D0 on the XIAO, which maps to "
    "GPIO1 in the ESP32's numbering — hence #define PIR_PIN 1 in the code."
)

bullet("Detection cone: roughly 110° in front of the sensor, up to about 5–7 metres.")
bullet("It cannot see. It has no idea what caused the heat shift — a person, a cat, a warm air vent, sunlight moving across a floor. It just says 'something warm moved.'")
bullet("No I2C address, no SPI clock, no data protocol. One wire, two states. Simplest possible sensor.")

note(
    "The Grove PIR (and most PIR modules) has a built-in hold time — after detecting motion "
    "the output stays HIGH for a fixed duration even if the motion has stopped. Many modules "
    "ship with this set to 3–8 seconds. Some have a tiny potentiometer on the PCB to adjust it."
)

doc.add_paragraph()

# ── Section 2 ────────────────────────────────────────────────────────────────

h1("2. How PIR connects to the camera")

body(
    "The PIR and the OV2640 camera are completely independent pieces of hardware. They do "
    "not talk to each other. The ESP32 chip in the middle talks to both:"
)

code("Grove PIR  ──signal wire──►  GPIO1 (D0)  ──ESP32 reads──►  your code decides")
code("                                                               │")
code("OV2640 camera ──parallel bus──► ESP32 ◄──esp_camera_fb_get()─┘")

body(
    "The camera is always running. From the moment esp_camera_init() is called, the OV2640 "
    "continuously captures frames and stores them in a frame buffer in PSRAM. It does this "
    "indefinitely — like a security camera always recording but the tape keeps overwriting itself."
)

body(
    "esp_camera_fb_get() is your way of saying 'give me the latest frame right now.' "
    "The camera does not take a photo on command — it hands you whatever it currently has "
    "buffered. Nothing gets stored unless your code explicitly calls fb_get() and copies "
    "the data somewhere. There is NO default auto-capture, no timer, no periodic save in "
    "the esp_camera library or OV2640 hardware."
)

doc.add_paragraph()

# ── Section 3 ────────────────────────────────────────────────────────────────

h1("3. How a person walking by gets captured")

body("Here is the sequence of events for one motion event, step by step:")

h3("a) Person enters the PIR's detection cone")
body(
    "The moment the person walks into the ~110° cone and their body heat shifts relative "
    "to the background, the PIR output goes HIGH."
)

h3("b) The ESP32 notices on the next loop() iteration")
body(
    "loop() runs thousands of times per second. Each iteration calls digitalRead(PIR_PIN). "
    "The code compares the current reading to lastPirState:"
)

code("bool pirNow = digitalRead(PIR_PIN);")
code("")
code("if (pirNow && !lastPirState) {")
code("    // Rising edge — runs ONCE at the moment motion starts")
code("    // If you skip this check, you capture on every loop() while PIR is HIGH")
code("}")
code("")
code("lastPirState = pirNow;  // remember for next iteration")

h3("c) The stale frame is discarded")
body(
    "The camera has been running continuously since boot. The frame currently in the buffer "
    "might be up to 66ms old. Discarding it forces the next fb_get() to wait for a genuinely "
    "fresh exposure taken after the PIR fired."
)

code("camera_fb_t* stale = esp_camera_fb_get();")
code("if (stale) esp_camera_fb_return(stale);")

h3("d) A fresh JPEG is captured")
code("uint32_t t0 = millis();")
code("camera_fb_t* fb = esp_camera_fb_get();")
code("uint32_t latency = millis() - t0;")
body(
    "fb->buf is a pointer to raw JPEG bytes. fb->len is how many bytes. "
    "This gets copied into PSRAM and served at /motion."
)

h3("e) PIR goes LOW")
body(
    "After the person leaves the detection cone, or the hold time expires, the PIR output "
    "drops LOW. The falling edge triggers a 2-second countdown before updating the baseline."
)

doc.add_paragraph()

# ── Section 4 ────────────────────────────────────────────────────────────────

h1("4. In-frame vs off-camera — what it means and how it works")

body(
    "The PIR's detection cone (~110°) is significantly wider than the OV2640's field of view "
    "(~66° for VGA). A person can absolutely trigger the PIR while standing completely outside "
    "what the camera can see. 'In frame' answers: was this person actually visible in the photo?"
)

h3("The trick: JPEG file size")
body(
    "JPEG compresses images by discarding detail. An empty room compresses extremely well — "
    "lots of uniform wall, floor, ceiling — producing a small file. A person adds texture, "
    "edges, and detail that JPEG cannot compress away, producing a noticeably larger file."
)

body("The code compares the captured frame's size against a learned baseline of the empty scene:")

code("float ratio    = (float)fb->len / baselineJpegLen;")
code("float absDelta = (float)fb->len - baselineJpegLen;")
code("bool  inFrame  = (ratio > 1.10f) && (absDelta > 5120.0f);")

body("Both conditions must be true:")
bullet("ratio > 1.10 — the frame must be at least 10% larger than a quiet scene.")
bullet("absDelta > 5120 — the frame must also be at least 5 KB larger in absolute terms. This prevents a noisy background from falsely triggering the in-frame flag.")

body(
    "If someone walks past the PIR but not the camera, the captured JPEG is roughly "
    "baseline size — inFrame = false. If they are visible, the JPEG is measurably bigger — "
    "inFrame = true."
)

doc.add_paragraph()

# ── Section 5 ────────────────────────────────────────────────────────────────

h1("5. The baseline and how it stays current")

body(
    "The baseline is the JPEG size of an empty quiet scene. It needs to adapt over time "
    "because lighting changes — morning vs afternoon, a lamp turning on — change what an "
    "'empty' room looks like to the camera."
)

body("The baseline uses an Exponential Moving Average (EMA):")

code("baselineJpegLen = 0.7f * baselineJpegLen + 0.3f * newSample;")

body(
    "It updates 2 seconds after the PIR goes quiet, giving the scene time to settle back "
    "to empty. The 70/30 split means it adapts slowly — a single anomalous frame cannot "
    "swing it far, but over many events it tracks gradual lighting changes."
)

note(
    "In the deep sleep version, the baseline is stored in RTC RAM using RTC_DATA_ATTR. "
    "This is a tiny 8KB memory region that stays powered during deep sleep, so the "
    "baseline survives across sleep cycles without needing to re-learn from scratch."
)

doc.add_paragraph()

# ── Section 6 ────────────────────────────────────────────────────────────────

h1("6. The 5-second photo mystery — diagnosed")

body(
    "When experimenting you observed photos being captured every ~5 seconds even with "
    "nothing apparently moving. Your theory was that the camera has a default refresh "
    "behaviour. That is incorrect — here is what actually happened:"
)

h3("The camera has no auto-capture")
body(
    "The OV2640 and the esp_camera library have no such behaviour. esp_camera_fb_get() "
    "is the ONLY way a frame leaves the camera's internal buffer. Nothing is stored "
    "automatically. The camera just keeps overwriting its own PSRAM buffer in a loop "
    "until you ask for a frame."
)

h3("Most likely cause A — PIR hold time")
body(
    "The Grove PIR (and most PIR modules using the HC-SR501 chip) has a built-in output "
    "hold time. After detecting motion, the output stays HIGH for a fixed duration — "
    "commonly 3–8 seconds — even if you have already moved away. Some PIR sensors also "
    "have a 'retriggerable' mode where any IR shift resets the hold timer."
)
body(
    "If something keeps triggering the PIR (a warm air vent, sunlight angle shifting, "
    "even the residual heat from your body after you have moved away), the PIR cycles "
    "HIGH → LOW → HIGH on roughly a 5-second rhythm. Combined with code that captures "
    "on every rising edge, you get photos every ~5 seconds."
)
body(
    "To diagnose: open Serial monitor and print a timestamp every time the PIR pin changes "
    "state. If you see it cycling on a regular interval with nobody in front of it, "
    "the hold time is the culprit."
)

h3("Most likely cause B — missing edge detection")
body("If your experimental code looked like this:")
code("void loop() {")
code("    if (digitalRead(PIR_PIN)) {")
code("        takePhoto();     // runs every loop() while PIR is HIGH")
code("    }")
code("    delay(5000);         // wait 5 seconds, check again")
code("}")
body(
    "This captures a photo, waits 5 seconds, then checks again. If the PIR hold time "
    "is still active, it captures again. You get photos at exactly 5-second intervals. "
    "The fix is the lastPirState edge detection — only capture on the LOW→HIGH transition."
)

h3("Possible cause C — web page auto-refresh")
body(
    "If the page you were checking had a <meta http-equiv='refresh' content='5'> tag "
    "(the /stats and /motions pages in main.cpp both do), the browser reloads every 5 "
    "seconds. If your /photo handler re-captured on every request rather than serving "
    "the last stored capture, you would see a new photo every 5 seconds from browser "
    "reloads alone — with no actual motion."
)

doc.add_paragraph()

# ── Section 7 ────────────────────────────────────────────────────────────────

h1("7. Battery saving — how deep sleep works")

body(
    "The ESP32S3 supports deep sleep: it shuts down almost everything (CPU, WiFi, camera) "
    "and draws around 14 µA instead of the usual 200–300 mA. GPIO1 (your PIR pin) is an "
    "RTC GPIO on the ESP32S3, which means it can be configured as a wake source."
)

h3("The tradeoff you have to accept")
body(
    "You cannot have deep sleep and a live stream at the same time. WiFi drops the moment "
    "you sleep. The camera de-initialises. The web server stops. While the board is asleep, "
    "http://10.10.10.155 is unreachable."
)

bullet("Current main.cpp (streaming): 200–300 mA, live stream available, motion capture works")
bullet("Deep sleep + PIR wake: ~14 µA sleeping, no stream, photo only on motion")

h3("What happens on each wake cycle")
body("Every time the PIR fires:")
bullet("Chip wakes (~1 ms)")
bullet("setup() runs again — camera init (~300 ms), WiFi reconnect (~3–8 seconds)")
bullet("Capture JPEG")
bullet("Serve it at /motion for 30 seconds, or POST to a server, or save to SD card")
bullet("esp_deep_sleep_start() — back to 14 µA")

h3("The code structure changes completely")
body(
    "With deep sleep there is no loop() in the traditional sense. setup() runs fresh "
    "on every wake. esp_sleep_get_wakeup_cause() tells you what triggered the wake:"
)

code("void setup() {")
code("    esp_sleep_wakeup_cause_t cause = esp_sleep_get_wakeup_cause();")
code("")
code("    if (cause == ESP_SLEEP_WAKEUP_EXT0) {")
code("        // PIR fired — init camera, capture, transmit")
code("    } else {")
code("        // First boot — take baseline, then sleep")
code("    }")
code("")
code("    esp_sleep_enable_ext0_wakeup(GPIO_NUM_1, 1); // wake on GPIO1 HIGH")
code("    esp_deep_sleep_start();  // never returns")
code("}")
code("")
code("void loop() {}  // never runs")

h3("Middle-ground option — light sleep")
body(
    "If you want some power saving but still want the web server to be reachable, "
    "use light sleep instead. The CPU pauses, WiFi maintains its connection with the "
    "router, and the camera stays initialised. Wake time is microseconds. Power drops "
    "to roughly 2–20 mA — not as low as deep sleep's 14 µA, but far better than 250 mA."
)

code("gpio_wakeup_enable(GPIO_NUM_1, GPIO_INTR_HIGH_LEVEL);")
code("esp_sleep_enable_gpio_wakeup();")
code("esp_light_sleep_start();  // pauses here until PIR fires, then continues")

doc.add_paragraph()

# ── Section 8 ────────────────────────────────────────────────────────────────

h1("8. The two-file approach")

body("You now have two separate firmware files, each doing one job well:")

h3("src/main.cpp — streaming + motion capture (always-on)")
bullet("Connects to WiFi on boot, stays connected")
bullet("HTTP server at port 80: /, /stream, /photo, /motion, /motions, /stats, /restart")
bullet("RTSP server at port 8554 for VLC")
bullet("PIR triggers in-PSRAM capture, ring buffer of last 3 events")
bullet("In-frame detection via JPEG size vs EMA baseline")
bullet("Power: ~200–300 mA continuously")

h3("practice/deep_sleep_motion.cpp — deep sleep (battery efficient)")
bullet("Sleeps at ~14 µA between events")
bullet("PIR fires → chip wakes → camera init → capture → WiFi → serve photo 30s → sleep")
bullet("Baseline stored in RTC RAM, survives sleep cycles")
bullet("In-frame detection still works — skip WiFi entirely if off-camera (Stage 5)")
bullet("No live stream. /motion only reachable for 30 seconds after a trigger.")
bullet("Power: ~14 µA sleeping, ~250 mA for ~10 seconds per motion event")

note(
    "To switch between them: copy the desired file's content into src/main.cpp "
    "(back up the original first), then build and upload via PlatformIO."
)

doc.add_paragraph()

# ── Section 9 ────────────────────────────────────────────────────────────────

h1("9. Learning roadmap — from scratch")

body(
    "If you are building the streaming version from scratch, work through these stages "
    "in order. Each one gives you something real and working before adding the next layer."
)

bullet("Stage 1", "Camera init + take one photo. Print its byte size to Serial. Works when you see a number like 14382 bytes.")
bullet("Stage 2", "Connect to WiFi, serve that photo at http://<ip>/photo. Works when a browser shows the photo.")
bullet("Stage 3", "Read PIR pin in loop(), print 'Motion!' / 'Quiet.' on state change. Works when waving your hand prints messages.")
bullet("Stage 4", "On PIR rising edge, capture a JPEG and serve it at /motion. Works when waving then opening /motion shows what it saw.")
bullet("Stage 5", "Store last 3 captures in a ring buffer. Serve them at /motion?n=0, ?n=1, ?n=2.")
bullet("Stage 6", "Track a quiet-scene baseline. Flag whether motion was in-frame or off-camera.")
bullet("Stage 7", "Serve a live MJPEG stream at /stream — continuous JPEG frames with a multipart boundary.")
bullet("Stage 8", "Add the full HTTP router. This is main.cpp — all stages combined.")

body(
    "For the deep sleep version, the stages are in practice/deep_sleep_motion.cpp: "
    "start with timed wake (timer instead of PIR), then PIR wake, then WiFi + serve, "
    "then RTC RAM baseline, then skip WiFi for off-camera events."
)

doc.add_paragraph()
divider()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run(
    "Hardware: Seeed XIAO ESP32S3 Sense + OV2640 + Grove PIR on D0 | "
    "Network: SenSen2 | Camera IP: 10.10.10.155 | RTSP: rtsp://10.10.10.155:8554/mjpeg/1"
).font.size = Pt(8)

# ── Save ─────────────────────────────────────────────────────────────────────

out = r"c:\Users\Gobind\Desktop\SENSEN WORK\s32\esp32-camera\practice\ESP32_PIR_Notes.docx"
doc.save(out)
print(f"Saved: {out}")
