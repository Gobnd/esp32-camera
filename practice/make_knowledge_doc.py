from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Style helpers ─────────────────────────────────────────────────────────────

DARK   = RGBColor(0x1a, 0x1a, 0x2e)
MID    = RGBColor(0x16, 0x21, 0x3e)
GREEN  = RGBColor(0x1a, 0x6b, 0x1a)
GREY   = RGBColor(0x55, 0x55, 0x55)
RED    = RGBColor(0x8b, 0x00, 0x00)
BLUE   = RGBColor(0x00, 0x44, 0x88)

def shade_para(p, hex_fill="F0F0F0"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = DARK
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = MID
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    return p

def body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def callout(label, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GREY
    shade_para(p, "F5F5F5")
    return p

def warn(text):  return callout("!", text, RED)
def tip(text):   return callout("→", text, BLUE)
def ahead(text): return callout("Plan ahead:", text, RGBColor(0x00, 0x66, 0x33))

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = GREEN
    shade_para(p)
    return p

def bullet(label, text=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    if text:
        rb = p.add_run(label + ": ")
        rb.bold = True; rb.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(label).font.size = Pt(11)

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(t.rows[0].cells):
        h.text = headers[i]
        for para in h.paragraphs:
            for run in para.runs:
                run.bold = True; run.font.size = Pt(10)
        tcPr = h._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'D0D8E8')
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(t.rows[ri+1].cells):
            val.text = str(row[ci])
            for para in val.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph("─" * 74)
    p.runs[0].font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    p.runs[0].font.size = Pt(9)

def page_break():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)

# ── Cover ─────────────────────────────────────────────────────────────────────

doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("ESP32-S3 Camera System")
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = DARK

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.add_run("How it works · Why it's built this way · What to do next").font.size = Pt(13)

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.add_run(f"Generated {datetime.date.today().strftime('%d %b %Y')}").font.size = Pt(10)

doc.add_paragraph()
divider()
doc.add_paragraph()

# ── Part labels ───────────────────────────────────────────────────────────────

def part(num, title, subtitle):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"PART {num}\n")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = GREY
    r2 = p.add_run(title)
    r2.bold = True; r2.font.size = Pt(18); r2.font.color.rgb = DARK
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(subtitle).font.size = Pt(11)
    divider()

# ═════════════════════════════════════════════════════════════════════════════
part(1, "Mental Models", "How every piece works, why it's designed that way, and how the pieces connect")
# ═════════════════════════════════════════════════════════════════════════════

# ─────────────────────────────────────────────────────────────────────────────
h1("1. The system in one paragraph")
# ─────────────────────────────────────────────────────────────────────────────

body("You have a small microcontroller (ESP32-S3) with a camera glued to it (OV2640) and a motion sensor attached (Grove PIR). The microcontroller connects to your WiFi and runs two things at once: a web server that lets any browser watch a live video feed, and a background detector that watches the PIR sensor and saves a photo to memory whenever someone walks by. That photo is viewable on a separate page in the browser. A second firmware variant skips the web server entirely, sleeps at near-zero power, and only wakes up when the PIR fires.")

body("Everything else in this document is detail on top of that paragraph. When something confuses you, come back to this and ask: which of those two things am I actually dealing with?")

# ─────────────────────────────────────────────────────────────────────────────
h1("2. Hardware mental models")
# ─────────────────────────────────────────────────────────────────────────────

h2("The ESP32-S3 — think of it as two workers sharing one toolbox")
body("The ESP32-S3 has two CPU cores. Think of them as two workers. They share the same tools (WiFi radio, camera, RAM) but can do different things at the same time.")
body("In this project Core 0 handles the RTSP video stream and the WiFi radio. Core 1 handles everything else: the web server, the PIR sensor, the motion captures, the LED. When Core 1 is busy serving five simultaneous video streams, the PIR is still being checked — just not as frequently, because Core 1 is juggling more tasks.")
tip("This is why latency during heavy streaming is higher than latency with no streams. It's not the camera being slow — it's Core 1 being interrupted by HTTP tasks before it can call fb_get().")

h2("The OV2640 camera — think of it as a tap, not a camera")
body("A normal camera takes a photo when you press a button. The OV2640 does not work that way. From the moment you initialise it, it runs continuously — producing a new JPEG frame every 40 ms (at 25 fps) and overwriting its internal buffer. It never stops. There is no trigger, no on-demand capture, no button.")
body("Your code does not take photos. Your code dips a cup into a running tap. esp_camera_fb_get() is the cup. If you call it, you get whatever is flowing right now. If you don't call it, the frames just disappear.")
tip("This is why there is no default 5-second capture behaviour in the camera itself. The camera is always producing. Something in your code was calling fb_get() every 5 seconds — the camera was not deciding to take a photo.")

h2("The PIR sensor — think of it as a light switch with a timer")
body("The PIR sensor does not see. It detects changes in infrared heat. When something warm moves across its field, it flips a wire from 0V to 3.3V. That is the only thing it does. It does not tell you what moved, how fast, or where. It just flips the wire.")
body("The critical detail: once it flips HIGH, it holds HIGH for a hardware-set duration (3–8 seconds on your module). This is called the hold time. It is set by a physical component on the sensor board, not by your firmware. When the hold time expires, the wire goes back to 0V. If the subject is still nearby and any heat movement occurs, it flips HIGH again immediately.")
warn("The hold time is why you see repeated captures 3–8 seconds apart with nothing obviously moving. The PIR is cycling through its hold time in response to residual heat, not the camera deciding to take a photo.")

h2("PSRAM — think of it as a big slow warehouse next to a small fast desk")
body("The ESP32-S3 has 512 KB of fast internal RAM (the desk) and 8 MB of external PSRAM (the warehouse). Camera frame buffers, motion capture copies, and the WiFi stack all live in the warehouse. The desk is for code execution and small variables.")
body("ps_malloc() allocates from the warehouse. malloc() allocates from the desk. If you use the wrong one for large allocations, the desk fills up and the board crashes. All motion capture buffers in this project use ps_malloc correctly.")
tip("The /stats page shows Free RAM which is the desk only. A PSRAM leak from failed free() calls would not appear there. No PSRAM leak test currently exists.")

# ─────────────────────────────────────────────────────────────────────────────
h1("3. The camera pipeline — from sensor to browser")
# ─────────────────────────────────────────────────────────────────────────────

h2("The double buffer and CAMERA_GRAB_LATEST")
body("The camera uses two PSRAM buffers — call them A and B. The camera writes to A while you read from B, then swaps. With CAMERA_GRAB_LATEST, if nobody is reading, the camera just keeps overwriting. Frames you never asked for disappear silently. This is deliberate — you always want the latest frame, not a queue of old ones.")
body("The consequence: if two consumers both call fb_get() at the same time, one of them has to wait. Consumer 1 holds buffer A. Consumer 2 calls fb_get() and gets buffer B. Now the camera has no free buffer to write into. It stalls. Consumer 2's call blocks until Consumer 1 returns buffer A.")
tip("With 5 HTTP stream clients active plus a PIR capture in progress, you have 6 things competing for 2 buffers. The motion capture's latency number includes this wait time — it is not purely camera speed.")

h2("Why the stale frame gets discarded on PIR trigger")
body("When the PIR fires, the buffer already contains a frame from before the trigger — maybe 40 ms before, maybe 80 ms before. That frame shows the empty corridor before the person entered. The firmware discards it and asks for the next one:")
code("{ camera_fb_t* s = esp_camera_fb_get(); if (s) esp_camera_fb_return(s); }")
code("camera_fb_t* fb = esp_camera_fb_get();  // this one is fresh")
body("Without the discard, every motion capture would show the scene from just before the person arrived — the wrong moment. With it, the capture shows the scene from just after the PIR fired.")

h2("JPEG size as a detection signal")
body("Here is the clever part of this design: the firmware does not do any image processing or AI. It uses JPEG file size as a proxy for 'is there something in frame.'")
body("JPEG compression works by finding patterns and redundancy. An empty room — mostly uniform wall, floor, ceiling — has enormous redundancy and compresses down to a small file. A person in frame breaks up those patterns with texture, clothing edges, and face detail. The same scene with a person in it produces a noticeably larger JPEG file.")
body("The firmware measures two things on every PIR-triggered capture:")
code("float ratio    = (float)fb->len / baselineJpegLen;   // e.g. 1.35 = 35% bigger")
code("float absDelta = (float)fb->len - baselineJpegLen;   // e.g. 6400 = 6.4 KB bigger")
code("bool  inFrame  = (ratio > 1.10f) && (absDelta > 5120.0f);")
body("Both conditions must be true. The ratio check catches proportional changes. The absolute check prevents a slightly noisy background from triggering in-frame just because 10% of a small baseline is an even smaller number.")

h2("The baseline — why it needs to adapt")
body("For the JPEG size comparison to work, the firmware needs to know what the empty scene looks like. That is the baseline. But lighting changes throughout the day — morning sun through a window makes the scene brighter and more detailed, which makes the empty-room JPEG larger. If the baseline never updated, morning captures would all look in-frame because the room itself got more complex.")
body("The baseline uses an Exponential Moving Average (EMA). It updates 2 seconds after the PIR goes quiet:")
code("baselineJpegLen = 0.7f * baselineJpegLen + 0.3f * newSample;")
body("The 70/30 split means any single sample can shift the baseline by at most 30%. A single unusual frame cannot corrupt it. But over many events it tracks gradual changes. Starting value: 20000 bytes (a conservative underestimate that biases toward more in-frame detections early on).")
warn("If the PIR goes LOW while the subject is still partially visible to the camera, the baseline update samples an occupied scene. Over multiple events the baseline drifts upward, making in-frame detection less sensitive. This is the most subtle ongoing failure mode in the firmware.")

# ─────────────────────────────────────────────────────────────────────────────
h1("4. The PIR pipeline — from heat to capture")
# ─────────────────────────────────────────────────────────────────────────────

h2("Why polling works here (and when it would not)")
body("The PIR is read with digitalRead() in every loop() iteration — no interrupt, no RTOS event. This works because the PIR hold time is 3–8 seconds. Even if loop() is delayed by 10 ms handling an HTTP request, a 3-second HIGH signal will still be caught. For a sensor with a 10 ms pulse (like a door switch or encoder), polling would miss events entirely and you would need an interrupt.")
tip("If you ever swap the Grove PIR for a faster sensor, add attachInterrupt(PIR_PIN, handler, RISING) and set a flag in the handler. Do not put heavy work in the interrupt itself — just set a flag and handle it in loop().")

h2("Edge detection — the key pattern")
body("The edge detection pattern is the most important piece of code for making PIR capture work correctly. Read it carefully:")
code("bool pirNow = digitalRead(PIR_PIN);")
code("")
code("if (pirNow && !lastPirState) {   // LOW→HIGH transition: fires exactly once")
code("    // capture here")
code("}")
code("")
code("lastPirState = pirNow;           // update EVERY iteration, not just on change")
body("pirNow is HIGH for the entire hold time — potentially thousands of loop() iterations. The condition pirNow && !lastPirState is only true for the single iteration where it transitions from LOW to HIGH. That is why you get one capture per motion event, not thousands.")
warn("lastPirState must be updated on every loop() iteration, not just inside the if block. If you move it inside the if, the condition is always true when the pin is HIGH and you are back to capturing every iteration.")

h2("The missing cooldown — the unfixed bug")
body("main.cpp correctly detects the rising edge once per transition. But it does not enforce a minimum time between captures. If the PIR hold time expires, the output goes LOW (falling edge), and then something triggers it again within a second, a new rising edge fires and another capture happens. In a warm room with slow-moving subjects this can produce captures every 5–8 seconds indefinitely.")
body("The fix is three lines:")
code("static unsigned long lastCapMs = 0;")
code("if (pirNow && !lastPirState && (millis() - lastCapMs > 5000)) {")
code("    lastCapMs = millis();")
code("    // capture block")
code("}")
ahead("This is not in main.cpp yet. Add it before you do any real-world deployment testing or T1 results will be confusing.")

# ─────────────────────────────────────────────────────────────────────────────
h1("5. The firmware architecture — why two cores matter")
# ─────────────────────────────────────────────────────────────────────────────

h2("Core assignment and what it means practically")
body("Core 0 runs the WiFi radio (lwIP TCP stack) and the RTSP streaming task. Core 1 runs Arduino's loop() and every HTTP connection handler. This split is deliberate: WiFi interrupt handling and RTSP streaming happen on Core 0 without competing with your motion detection code on Core 1.")
body("The problem is that Core 1 is shared between loop() (where PIR polling happens) and all HTTP handler tasks (where streaming happens). FreeRTOS gives each task a time slice. With 5 stream clients, Core 1 has 6 things competing: 5 stream handlers and loop(). Loop() still runs but gets interrupted more frequently.")

h2("The streamFps race — why FPS numbers lie under multi-client load")
body("Every HTTP stream handler writes to the same global variable:")
code("streamFps = fpsFrameCount * 1000.0f / (float)(now - fpsWindowStart);")
body("With 3 clients running, all three tasks write to streamFps. The value in /stats is whichever task happened to write it most recently. It is not total throughput. It is not an average. It is one task's rate from an arbitrary recent moment.")
body("The test doc's T5 FPS numbers (27.8 → 12.4 → 7.3 as clients increase) look like the board is degrading badly. Some of that drop is real (the board is handling more load). Some of it is this race condition making the reported number unreliable.")
tip("The FPS number that IS reliable: measure it on the client side in Python by counting --frame boundary markers yourself. That is what res_fps_test.py does correctly.")

h2("The mutex — what it protects and what it costs")
body("The motionSlots ring buffer is written by loop() and read by HTTP handler tasks simultaneously. Without protection, a reader could see a half-written slot — buf pointing to freed memory, len reflecting the old value, ms reflecting the new one. The motionMutex prevents this.")
body("The cost: loop() currently holds the mutex for the entire duration of ps_malloc + memcpy. At VGA quality 20 that is ~30 KB copied under lock. At SVGA quality 10 that could be ~100 KB. During that time, any HTTP request to /motion or /motions blocks for the copy duration.")
ahead("If you switch to SVGA or HD for captures (better face detail), the mutex hold time grows significantly. Refactor to copy outside the mutex and only take it briefly to swap the pointer before deploying at higher resolutions.")

# ─────────────────────────────────────────────────────────────────────────────
h1("6. The deep sleep firmware — a different design philosophy")
# ─────────────────────────────────────────────────────────────────────────────

h2("The fundamental shift: setup() is the program")
body("In the streaming firmware, setup() runs once at boot and loop() runs forever. In the deep sleep firmware, there is no meaningful loop(). setup() runs on every wake — whether that is first boot or PIR trigger — and always ends by going back to sleep. The 'program' is one wake cycle:")
code("boot → setup() → [do one thing] → esp_deep_sleep_start() → [PIR fires] → setup() → ...")
body("esp_deep_sleep_start() never returns. The chip powers down. When the PIR fires, the chip boots from scratch — setup() runs again with no memory of what happened before, except for variables marked RTC_DATA_ATTR which survive in the 8 KB RTC slow memory.")

h2("What survives sleep and what does not")
table(
    ["Memory type", "Survives deep sleep?", "What lives there"],
    [
        ["RTC slow memory (8 KB)", "YES", "RTC_DATA_ATTR variables: bootCount, baselineJpegLen"],
        ["RTC fast memory (8 KB)", "YES", "Available, not currently used"],
        ["Internal SRAM (512 KB)",  "NO",  "All normal variables, stack, heap"],
        ["PSRAM (8 MB)",            "NO",  "Camera frame buffers, ps_malloc allocations"],
        ["Flash",                   "YES", "Firmware, constants — always survives"],
    ],
    col_widths=[2.2, 1.5, 2.7]
)
body("This is why the baseline can survive across sleep cycles (it is RTC_DATA_ATTR) but the captured JPEG cannot (it was in PSRAM). On every wake, the firmware must re-initialise the camera, re-connect to WiFi, and re-allocate everything from scratch.")

h2("The wake-up sequence and its timing")
body("Here is the full timeline for one motion event in deep sleep mode:")
table(
    ["Step", "Who does it", "Typical time"],
    [
        ["PIR fires, GPIO1 goes HIGH",    "Hardware",    "0 ms"],
        ["RTC domain wakes CPU",           "Hardware",    "~1 ms"],
        ["setup() runs, Serial init",      "Firmware",    "~500 ms"],
        ["esp_camera_init()",              "Firmware",    "~300 ms"],
        ["Discard stale frame + capture",  "Firmware",    "~40–80 ms"],
        ["WiFi connect",                   "Firmware",    "3,000–8,000 ms"],
        ["Serve /motion for 30 s",         "Firmware",    "30,000 ms"],
        ["WiFi disconnect + sleep",        "Firmware",    "~200 ms"],
    ],
    col_widths=[2.8, 1.5, 1.5]
)
warn("The person who triggered the PIR walked past 4–9 seconds before the photo is even captured. At a brisk walking pace that is 4–8 metres of movement since the trigger. This is the fundamental tradeoff of deep sleep. The captured frame shows the scene as it was 4–9 seconds after the trigger — not the moment of detection.")
ahead("If capture latency matters more than battery life, use light sleep instead. Light sleep keeps WiFi associated and the camera initialised. Wake time drops to under 1 ms. Power draw is ~2–20 mA instead of ~14 µA.")

# ─────────────────────────────────────────────────────────────────────────────
h1("7. What breaks and why — the connected failure modes")
# ─────────────────────────────────────────────────────────────────────────────

body("The failures in this system are not random. They connect to each other through a common theme: the firmware has no concept of how long a motion event lasts. It knows the moment the PIR fired. It does not know whether the person is still there, how long they have been there, or whether the scene has settled back to empty.")

h2("The chain: hold time → re-trigger → repeated captures → baseline drift")
body("Step 1: Person walks past. PIR fires. Capture taken.")
body("Step 2: PIR hold time runs (3–8 s). Person is gone but PIR stays HIGH.")
body("Step 3: Hold time expires. PIR goes LOW. pirFellMs is set. 2-second baseline countdown begins.")
body("Step 4: Person's body heat is still radiating from the floor/wall. PIR retriggers. New rising edge. New capture fires. pirFellMs is reset.")
body("Step 5: The 2-second window never completes because the PIR keeps cycling. Baseline is never updated.")
body("Step 6: Over many events, the baseline slowly drifts toward the occupied-scene value because it occasionally updates when the person is still partially in frame.")
body("Step 7: In-frame threshold drifts upward. Future events increasingly report off-camera even when someone is clearly visible.")
tip("All of this is fixed by one change: a 5-second minimum capture interval. It breaks the hold-time cycle at step 4 and gives the baseline update a clean window to run.")

h2("The streamFps number and what tests it corrupts")
body("Because streamFps is a race condition between concurrent handlers, any test that reads it under load is reading noise. This affects: T4 (FPS column for multi-stream rows), T5 (FPS column entirely), T6b (FPS column), wifi_test.py (both FPS readings). The FPS numbers in those tests are directionally right — more clients = lower number — but the specific values are not reproducible and should not be treated as hard data.")

# ═════════════════════════════════════════════════════════════════════════════
page_break()
part(2, "Planning Ahead", "What to think about before you do the next thing")
# ═════════════════════════════════════════════════════════════════════════════

# ─────────────────────────────────────────────────────────────────────────────
h1("8. Before you run the pending tests")
# ─────────────────────────────────────────────────────────────────────────────

h2("T1 — Motion detection")
body("The test polls /stats every second, detects when the captures counter increments, then fetches /motions to get the latency and in-frame badge. There are three things that will make your T1 results unreliable if you do not account for them:")
bullet("Close all browser tabs first", "An active stream client holds a frame buffer. PIR capture has to wait for it. Your latency numbers will be inflated by buffer contention, not camera speed. Run T1 with no other clients.")
bullet("The captures counter increments on every rising edge", "Without the cooldown fix, one corridor walk-through can produce 3–4 captures as the PIR cycles through its hold time. The ring buffer fills with near-duplicate frames. Your T1 table will show 4 events from one walk-through.")
bullet("The test timestamp is Python time, not firmware time", "The test detects events by polling every second. An event that happened at t=4.1s appears in the test at t=5.0s. The latency column (from /motions) is accurate — it is measured by the firmware. The event timestamp is not.")
ahead("Add the 5-second cooldown to main.cpp before running T1. Otherwise you are measuring the PIR hold time behaviour, not motion detection performance.")

h2("T2 — WiFi range")
body("wifi_test.py reads streamFps from /stats (the racy global). Under 2 clients, whichever stream handler last wrote it wins. The comparative FPS (1 viewer vs 2 viewers at the same location) is still useful — both readings are equally noisy, so the relative drop is meaningful even if the absolute numbers are not.")
bullet("Move the board on a battery bank between locations", "Reconnect USB for Serial if you want to monitor, or just use browser access to /stats")
bullet("Wait the full 8 seconds after opening each stream", "The FPS takes 1–2 seconds to stabilise. 8 seconds is more than enough.")
bullet("RSSI does not change between the 1-viewer and 2-viewer rows", "It is a property of the RF environment. Only compare RSSI across locations, not across viewer counts.")
ahead("T2 will show you the practical deployment range. Below −70 dBm FPS drops significantly. If your deployment location is marginal, that is when you need to think about a WiFi repeater or moving the router.")

h2("T1 deep sleep version — what to expect")
body("When you test deep_sleep_motion.cpp, the metrics are different from main.cpp T1. There is no captures counter. There is no /stats. The only observable outputs are Serial log lines and the 30-second /motion window. Before testing:")
bullet("Open Serial monitor at 115200 baud before flashing", "You need to see the first boot baseline capture and the wake cause on PIR trigger")
bullet("On first boot you will see 'First boot — capturing baseline'", "This is correct. The board takes one photo to set the baseline, then sleeps immediately with no WiFi.")
bullet("On PIR wake you will see the JPEG size, baseline, ratio, and in-frame verdict", "If ratio < 1.10 or absDelta < 5120, it will say off-camera even if you are standing in front of it — this means the baseline is too close to the occupied-scene size")
ahead("If you consistently get off-camera in deep sleep mode, physically cover the camera for 30 seconds after first boot to force a clean empty-scene baseline before your first real trigger.")

# ─────────────────────────────────────────────────────────────────────────────
h1("9. Before you change the firmware")
# ─────────────────────────────────────────────────────────────────────────────

h2("If you increase resolution for captures")
body("main.cpp uses VGA for both streaming and motion capture. If you want higher resolution captures (e.g. HD for face detail at longer range), you have two problems:")
bullet("The mutex hold time grows", "HD JPEG at quality 20 is ~80–120 KB. That memcpy under mutex blocks /motion requests for the full copy duration. Refactor to copy outside the mutex first.")
bullet("FPS drops below 20", "HD gives 13.9 fps. Motion capture works fine at any fps — it only grabs one frame per event. But your live stream will visibly stutter.")
tip("One approach: stream at VGA for the live view but set a higher resolution specifically for captures. Call s->set_framesize() at the start of the PIR capture block and restore VGA afterward.")
ahead("If you do this, the baseline was learned at VGA. HD jpegs are naturally larger even for an empty room. Reset the baseline after switching capture resolution or in-frame detection will be wrong.")

h2("If you add a second camera")
body("Every file in this project has the IP address 10.10.10.155 hardcoded: main.cpp uses it nowhere (it reads the IP at runtime), but CLAUDE.md, all three Python test scripts, and the test document all hardcode it. If a second board gets a different IP from your router, every script needs updating.")
ahead("Before you get a second unit: change the test scripts to accept IP as a command-line argument instead of hardcoding it. One line change per script. Do it now before you have two boards and are editing six files under pressure.")

h2("If you want to add face recognition")
body("The current system captures and stores JPEGs. It does no recognition. Adding recognition means processing those frames somewhere — either on the ESP32 itself (very limited, ESP32 is not an ML accelerator) or on a separate computer that receives the stream.")
body("For a separate computer doing recognition:")
bullet("Use RTSP, not HTTP /stream", "RTSP is designed for this use case. HTTP MJPEG is browser-friendly but inefficient for ML pipelines. And only one RTSP client is supported — your recognition software should be the sole consumer.")
bullet("Use timestamps, not frame counts", "The T6b stability data shows FPS jittering between 5.5 and 9.2 fps under load. Frame counts are unreliable for timing. Always use millis() timestamps on the firmware side and wall-clock time on the receiving side.")
bullet("VGA is the right resolution", "150 cm face recognition range at 25 fps beats 300 cm at 14 fps for a door or corridor. At 14 fps a fast-walking person gives you fewer usable frames per transit.")
ahead("When you add recognition, the 'one RTSP client' limit becomes a real constraint. You cannot have VLC open for monitoring AND recognition software running simultaneously on RTSP. You need to choose, or add a second HTTP stream consumer and accept the FPS cost.")

h2("If you want to save captures to an SD card")
body("The XIAO ESP32S3 Sense has a microSD card slot on the expansion board. The firmware does not currently use it. Adding it would let the deep sleep firmware store captures locally without WiFi — removing the 4–8 second WiFi reconnect delay entirely.")
ahead("The SD card uses SPI. Check which SPI pins are used on the Sense expansion board before writing code — some SPI pins overlap with camera data lines on certain board revisions. Verify with the Seeed XIAO ESP32S3 Sense schematic before assuming any SPI pin is free.")

# ─────────────────────────────────────────────────────────────────────────────
h1("10. Things that will confuse you if you do not know them")
# ─────────────────────────────────────────────────────────────────────────────

bullet("The camera body reaching 62°C under load can cause PIR false triggers",
       "The PIR detects heat changes. If the camera is mounted close to the PIR sensor, the camera's own warmth can drift the PIR's reference level and cause spurious triggers as the board heats up from idle to streaming load. Mount PIR ≥5 cm from the camera body.")

bullet("esp_deep_sleep_start() in setup() means the board appears to do nothing",
       "On first boot of the deep sleep firmware, the board takes a baseline photo, prints one line to Serial, and immediately sleeps. If you are not watching Serial monitor, it looks like the upload failed or the board crashed. It did not — it is sleeping, waiting for PIR.")

bullet("Resetting the board does not reset the baseline in deep sleep mode",
       "baselineJpegLen is RTC_DATA_ATTR. A soft reset (pressing the reset button, or ESP.restart()) does NOT clear RTC RAM. The baseline persists. Only a full power cycle or holding the boot button clears it. If your baseline has drifted and in-frame detection is unreliable, power cycle the board.")

bullet("The RTSP stream resolution lags one reconnect behind",
       "When you change resolution via /res, the HTTP stream updates within 1–2 frames. The RTSP stream does not. The RTSP task picks up currentW and currentH only when a new client connects. VLC must disconnect and reconnect to see the new resolution.")

bullet("allow_redirects=False in the test scripts is intentional",
       "/res, /quality, and /fps all return HTTP 302 redirects to /. The firmware applies the setting before sending the 302. The test scripts use allow_redirects=False to stop requests from following the redirect — otherwise they would open a second connection to / and inflate the connection count. The setting IS applied even though the test does not follow the redirect.")

bullet("/stats FPS reads as 0 when no browser is watching",
       "streamFps is set to 0 in setup() and only updated while a stream handler is running. If you open /stats with no active stream client, FPS shows 0. This does not mean the camera is broken — it means no one is watching the stream right now.")

# ═════════════════════════════════════════════════════════════════════════════
page_break()
part(3, "Field Reference", "Look things up when you are at the board")
# ═════════════════════════════════════════════════════════════════════════════

# ─────────────────────────────────────────────────────────────────────────────
h1("11. Hardware quick reference")
# ─────────────────────────────────────────────────────────────────────────────

h2("Board and network")
table(
    ["Item", "Value"],
    [
        ["Board",       "Seeed XIAO ESP32S3 Sense (OV2640 built-in)"],
        ["MCU",         "ESP32-S3, dual-core LX7, 240 MHz"],
        ["RAM",         "512 KB internal + 8 MB PSRAM"],
        ["WiFi SSID",   "SenSen2"],
        ["Password",    "wongabongamcdonga"],
        ["Camera IP",   "10.10.10.155"],
        ["Web UI",      "http://10.10.10.155"],
        ["RTSP URL",    "rtsp://10.10.10.155:8554/mjpeg/1"],
        ["Serial baud", "115200"],
        ["Deep sleep",  "~14 µA"],
        ["Active load", "~200–310 mA (240 MHz)"],
        ["Temp limit",  "85°C (throttle). Normal streaming: 56–64°C"],
    ],
    col_widths=[1.8, 4.2]
)

h2("PIR wiring")
table(
    ["Wire colour", "Connect to"],
    [
        ["Red",    "3V3"],
        ["Black",  "GND"],
        ["Yellow", "D0 (GPIO1)"],
    ],
    col_widths=[1.5, 2.5]
)

h2("Camera pin assignments (do not change)")
table(
    ["Signal", "GPIO", "Signal", "GPIO", "Signal", "GPIO"],
    [
        ["XCLK",  "10", "PCLK", "13", "VSYNC", "38"],
        ["HREF",  "47", "SIOD", "40", "SIOC",  "39"],
        ["Y2/D0", "15", "Y3/D1","17", "Y4/D2", "18"],
        ["Y5/D3", "16", "Y6/D4","14", "Y7/D5", "12"],
        ["Y8/D6", "11", "Y9/D7","48", "PIR",    "1"],
    ],
    col_widths=[1.0, 0.7, 1.0, 0.7, 1.0, 0.7]
)

# ─────────────────────────────────────────────────────────────────────────────
h1("12. HTTP endpoints")
# ─────────────────────────────────────────────────────────────────────────────

table(
    ["Endpoint", "What it does"],
    [
        ["/",               "Web UI — live stream, resolution/quality/FPS controls"],
        ["/stream",         "Raw MJPEG stream — keep connection open, browser updates image each frame"],
        ["/photo",          "Single JPEG snapshot of current frame"],
        ["/motion?n=0",     "Most recent PIR capture as JPEG. n=1 second, n=2 oldest. 503 if none yet."],
        ["/motions",        "Gallery of last 3 captures. Auto-refreshes every 5 s."],
        ["/stats",          "All metrics: uptime, CPU, RAM, PSRAM, temp, WiFi, FPS, motion history. Refreshes every 3 s."],
        ["/res?id=vga",     "Set resolution. IDs: qvga, vga, svga, hd, sxga, uxga. Redirects to /."],
        ["/quality?q=20",   "Set JPEG quality 4–63. Lower = better quality, bigger file."],
        ["/fps?delay=0",    "Frame delay: 0=max, 33=~30fps, 66=~15fps, 200=~5fps."],
        ["/restart",        "Reboots the board in ~5 s."],
    ],
    col_widths=[2.0, 4.0]
)

# ─────────────────────────────────────────────────────────────────────────────
h1("13. Test results")
# ─────────────────────────────────────────────────────────────────────────────

h2("Resolution vs FPS — T3b (measured 2026-05-14)")
table(
    ["Resolution", "Dimensions", "FPS", "Temp", "Face clear at", "≥20fps?"],
    [
        ["qvga", "320×240",   "25.4", "61.3°C", "100 cm", "YES"],
        ["vga",  "640×480",   "24.8", "62.3°C", "150 cm", "YES — recommended"],
        ["svga", "800×600",   "13.1", "59.3°C", "—",      "NO"],
        ["hd",   "1280×720",  "13.9", "62.3°C", "300 cm", "NO"],
        ["sxga", "1280×1024", "10.6", "62.3°C", "320 cm", "NO"],
        ["uxga", "1600×1200", "7.8",  "63.3°C", "380 cm", "NO"],
    ],
    col_widths=[0.9, 1.1, 0.7, 0.9, 1.3, 1.5]
)

h2("Face recognition distance — T3 (measured 2026-05-14, fair indoor lighting)")
table(
    ["Resolution", "Clearly", "Barely", "Gone"],
    [
        ["240p 320×240",   "100 cm", "150 cm", "200 cm"],
        ["480p 640×480",   "150 cm", "200 cm", "250 cm"],
        ["720p 1280×720",  "300 cm", "350 cm", "400 cm"],
        ["SXGA 1280×1024", "320 cm", "360 cm", "400 cm"],
        ["UXGA 1600×1200", "380 cm", "450 cm", "500 cm"],
    ],
    col_widths=[1.8, 1.2, 1.2, 1.2]
)
tip("All distances drop 30–50% in dim lighting.")

h2("5-minute stability — T6b (measured 2026-05-14, 3 streams)")
table(
    ["Elapsed", "Alive", "FPS", "Temp", "Free RAM", "RSSI"],
    [
        ["30s",  "3/3", "9.2", "56.3°C", "190/323 KB", "−38 dBm"],
        ["60s",  "3/3", "6.5", "57.3°C", "183/323 KB", "−38 dBm"],
        ["120s", "3/3", "6.5", "59.3°C", "191/323 KB", "−39 dBm"],
        ["180s", "3/3", "6.5", "60.3°C", "182/323 KB", "−38 dBm"],
        ["240s", "3/3", "6.7", "61.3°C", "190/323 KB", "−39 dBm"],
        ["300s", "3/3", "7.0", "62.3°C", "191/323 KB", "−38 dBm"],
    ],
    col_widths=[0.8, 0.7, 0.6, 0.9, 1.3, 1.1]
)
body("PASS. Temp plateaued at 62°C. RAM stable (no heap leak). RSSI strong throughout.")

# ─────────────────────────────────────────────────────────────────────────────
h1("14. Running the test scripts")
# ─────────────────────────────────────────────────────────────────────────────

body("All scripts require main.cpp flashed. Run from the esp32-camera/ directory. Close all browser tabs first.")
table(
    ["Script", "Command", "What it runs", "When to use it"],
    [
        ["run_tests.py",    "python run_tests.py --motion",  "T4, T6, TEP, T5, T6b, T1",    "Full test pass. Add --motion only if PIR is wired."],
        ["res_fps_test.py", "python res_fps_test.py",        "FPS at all 6 resolutions",     "After any firmware change that could affect FPS. Run in isolation."],
        ["wifi_test.py",    'python wifi_test.py "location"',"RSSI + FPS at 1 and 2 viewers","Once per physical location with board on battery bank."],
    ],
    col_widths=[1.5, 2.2, 1.8, 2.0]
)
warn("res_fps_test.py will now correctly warn you if another client is streaming (bug fixed 2026-05-19). Do not proceed past that warning or FPS numbers will be wrong.")

# ─────────────────────────────────────────────────────────────────────────────
h1("15. File map")
# ─────────────────────────────────────────────────────────────────────────────

table(
    ["File", "Status", "Purpose"],
    [
        ["src/main.cpp",                       "Active firmware",  "Streaming + always-on. All tests target this. Do not break."],
        ["practice/deep_sleep_motion.cpp",     "Copy to use",      "Deep sleep PIR-wake firmware. Copy to src/main.cpp to flash."],
        ["practice/pir_practice.cpp",          "Learning only",    "Staged roadmap for building the streaming firmware yourself."],
        ["run_tests.py",                        "Run directly",     "Full automated test suite."],
        ["res_fps_test.py",                     "Run directly",     "Resolution vs FPS. Pre-check now works (fixed 2026-05-19)."],
        ["wifi_test.py",                        "Run directly",     "WiFi range and FPS per location."],
        ["platformio.ini",                      "Build config",     "Board: seeed_xiao_esp32s3. Library: Micro-RTSP."],
    ],
    col_widths=[2.5, 1.3, 2.7]
)

doc.add_paragraph()
divider()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run(
    "Board: Seeed XIAO ESP32S3 Sense  ·  IP: 10.10.10.155  ·  "
    "Firmware date: 2026-05-14  ·  Doc generated: " + datetime.date.today().strftime('%Y-%m-%d')
).font.size = Pt(8)

out = r"c:\Users\Gobind\Desktop\SENSEN WORK\s32\esp32-camera\practice\ESP32_Camera_Knowledge.docx"
doc.save(out)
print(f"Saved: {out}")
